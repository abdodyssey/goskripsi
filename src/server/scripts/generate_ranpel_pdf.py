import sys
import json
import os
import requests
from docx import Document
from docx.shared import Inches
import subprocess
from io import BytesIO

import re
import html

def strip_html(text):
    if not text:
        return ""
    # Replace <p> and <br> with newlines to preserve spacing
    text = text.replace('</p>', '\n').replace('<br>', '\n').replace('<br/>', '\n').replace('<br />', '\n')
    # Remove all other HTML tags
    clean = re.compile('<.*?>')
    cleantext = re.sub(clean, '', text)
    # Unescape HTML entities (e.g. &nbsp;)
    return html.unescape(cleantext).strip()

def replace_text_in_runs(paragraph, placeholder, replacement):
    if placeholder in paragraph.text:
        for run in paragraph.runs:
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, replacement)

def generate_ranpel_pdf(json_path, template_path, output_pdf_path):
    try:
        with open(json_path, 'r') as f:
            data = json.load(f)
        doc = Document(template_path)

        # 1. Fill basic info (Only header info)
        for p in doc.paragraphs:
            if "Nama Mahasiswa (NIM)-" in p.text:
                p.text = f"Nama Mahasiswa (NIM)- {data.get('studentName', '')} ({data.get('studentNim', '')})"
            
            if "Palembang, " in p.text:
                p.text = f"Palembang, {data.get('tanggal', '')}"

        # 2. Fill content and remove instructions
        content_found = {
            'judul': False,
            'masalah': False,
            'solusi': False,
            'hasil': False,
            'data': False,
            'metode': False,
            'jurnal': False
        }

        # Helper to set paragraph text with indentation
        def set_indented_text(p, text):
            p.text = text
            # Indent to ~0.5 inches to match the list style
            p.paragraph_format.left_indent = Inches(0.5)

        for i, p in enumerate(doc.paragraphs):
            # Judul Penelitian
            if "Judul Penelitian" in p.text and not content_found['judul']:
                if i + 1 < len(doc.paragraphs):
                    set_indented_text(doc.paragraphs[i+1], strip_html(data.get('judulPenelitian', '')))
                content_found['judul'] = True
            
            # Masalah dan Penyebab
            if "Membahas tentang permasalahan" in p.text:
                set_indented_text(p, strip_html(data.get('masalahDanPenyebab', '')))
                content_found['masalah'] = True
            elif any(x in p.text for x in ["Jika penelitian berkaitan", "Sebaiknya dilengkapi dengan data awal"]):
                p.text = ""
            
            # Alternatif Solusi
            if "embahas cara-cara atau solusi" in p.text:
                set_indented_text(p, strip_html(data.get('alternatifSolusi', '')))
                content_found['solusi'] = True
            elif any(x in p.text for x in ["Tersebut.", "Memiliki dasar teori", "Sumber referensi diutamakan"]):
                p.text = ""
            
            # Hasil yang diharapkan
            if "Memuat target yang akan dicapai" in p.text:
                set_indented_text(p, strip_html(data.get('hasilYangDiharapkan', '')))
                content_found['hasil'] = True
            elif "Berupa hasil (output) dan dampak" in p.text:
                p.text = ""
            
            # Kebutuhan Data
            if "Data yang akan diolah" in p.text:
                set_indented_text(p, strip_html(data.get('kebutuhanData', '')))
                content_found['data'] = True
            elif any(x in p.text for x in ["Data awal, sebagai bahan", "Bahan penelitian, bahan yang akan diolah"]):
                p.text = ""
            
            # Metode Pelaksanaan
            if "Metode yang akan digunakan" in p.text:
                set_indented_text(p, strip_html(data.get('metodePenelitian', '')))
                content_found['metode'] = True
            elif any(x in p.text for x in ["Misal: Metode analisis", "metode perancangan", "metode pengujian"]):
                p.text = ""
                
            # Jurnal Referensi
            if "Penelitian-penelitian yang dijadikan rujukan" in p.text:
                set_indented_text(p, strip_html(data.get('jurnalReferensi', '')))
                content_found['jurnal'] = True
            elif any(x in p.text for x in ["Dimuat pada:", "prosiding", "nama perguruan tinggi"]):
                p.text = ""

        # 3. Handle Signatures with a Table for perfect alignment
        for i, p in enumerate(doc.paragraphs):
            if "Dosen PA," in p.text and "Penulis" in p.text:
                # 1. Clear the paragraphs in the signature section (P34 to P38 approx)
                # P34: Dosen PA, Penulis
                # P35-36: Empty
                # P37: Nama Dosen PA / Nama Mahasiswa
                # P38: NIP / NIM
                
                # We'll clear this paragraph and the next several until we hit Lampiran
                p.text = ""
                for j in range(1, 10): # Increase range to find Lampiran even if there are many empty lines
                    if i + j < len(doc.paragraphs):
                        target_p = doc.paragraphs[i+j]
                        if "Lampiran:" in target_p.text:
                            # Also reduce spacing above Lampiran
                            target_p.paragraph_format.space_before = 0
                            target_p.paragraph_format.space_after = 0
                            target_p.paragraph_format.line_spacing = 1.0
                            break
                        target_p.text = ""
                        target_p.paragraph_format.space_before = 0
                        target_p.paragraph_format.space_after = 0
                        target_p.paragraph_format.line_spacing = 0.7 # Extra tight

                # 2. Add a table for signatures with specific column widths for alignment
                table = doc.add_table(rows=4, cols=2)
                # Set table to be full width or at least wide enough
                table.autofit = False
                
                # A4 width is ~6.3 inches. Let's make first col ~3.5 and second ~2.8
                # This should push the second column to the right to align with "Palembang"
                table.columns[0].width = Inches(3.5)
                table.columns[1].width = Inches(2.5)
                
                # Helper to set cell alignment and text
                def set_cell(row, col, text, bold=False):
                    cell = table.cell(row, col)
                    cell.text = text
                    # Left align paragraph in cell (0 is WD_ALIGN_PARAGRAPH.LEFT)
                    cell.paragraphs[0].alignment = 0
                
                # Row 0: Titles
                set_cell(0, 0, "Dosen PA,")
                set_cell(0, 1, "Penulis")
                
                # Row 1: Signature Images / Space (Set height for enough space)
                table.rows[1].height = Inches(0.8)
                
                # Row 1: Signature Images
                if data.get('dosenPaSignatureUrl'):
                    try:
                        response = requests.get(data['dosenPaSignatureUrl'], timeout=5)
                        if response.status_code == 200:
                            img_stream = BytesIO(response.content)
                            cell = table.cell(1, 0)
                            cell.paragraphs[0].alignment = 0
                            run = cell.paragraphs[0].add_run()
                            run.add_picture(img_stream, height=Inches(0.8))
                    except: pass
                
                if data.get('studentSignatureUrl'):
                    try:
                        response = requests.get(data['studentSignatureUrl'], timeout=5)
                        if response.status_code == 200:
                            img_stream = BytesIO(response.content)
                            cell = table.cell(1, 1)
                            cell.paragraphs[0].alignment = 0
                            run = cell.paragraphs[0].add_run()
                            run.add_picture(img_stream, height=Inches(0.8))
                    except: pass
                
                # Row 2: Names
                set_cell(2, 0, data.get('dosenPaNama', ''))
                set_cell(2, 1, data.get('studentName', ''))
                
                # Row 3: NIP/NIM
                set_cell(3, 0, f"NIP. {data.get('dosenPaNip', '')}")
                set_cell(3, 1, f"NIM. {data.get('studentNim', '')}")
                
                # Close the gap even more by setting zero spacing on the last row
                table.cell(3, 0).paragraphs[0].paragraph_format.space_after = 0
                table.cell(3, 1).paragraphs[0].paragraph_format.space_after = 0
                table.cell(3, 0).paragraphs[0].paragraph_format.line_spacing = 1.0
                table.cell(3, 1).paragraphs[0].paragraph_format.line_spacing = 1.0
                
                # We need to move the table to where the paragraph was.
                # In python-docx, new tables are added at the end.
                # To insert it at a specific place, we have to use the underlying XML.
                # Or a simpler trick: find a way to place it.
                # Actually, doc.add_table appends.
                # To move it:
                p._element.addnext(table._element)
                break

        # Save temporary docx in the same directory as output
        output_dir = os.path.dirname(os.path.abspath(output_pdf_path)) or os.getcwd()
        temp_docx = os.path.join(output_dir, f"temp_{os.path.basename(output_pdf_path).replace('.pdf', '.docx')}")
        doc.save(temp_docx)

        # Convert to PDF using LibreOffice
        subprocess.run([
            'libreoffice', '--headless', '--convert-to', 'pdf', temp_docx, '--outdir', output_dir
        ], check=True, capture_output=True)

        # LibreOffice saves it as temp_xxx.pdf in output_dir
        generated_pdf = temp_docx.replace('.docx', '.pdf')
        if os.path.exists(generated_pdf):
            os.rename(generated_pdf, output_pdf_path)
        else:
            print(f"Error: Generated PDF not found at {generated_pdf}")
            sys.exit(1)
        
        # Cleanup
        if os.path.exists(temp_docx):
            os.remove(temp_docx)
            
        print(f"Successfully generated PDF: {output_pdf_path}")

    except Exception as e:
        print(f"Error: {e}")
        sys.exit(1)

if __name__ == "__main__":
    if len(sys.argv) < 4:
        print("Usage: python generate_ranpel_pdf.py <json_data> <template_path> <output_pdf_path>")
        sys.exit(1)
    generate_ranpel_pdf(sys.argv[1], sys.argv[2], sys.argv[3])
