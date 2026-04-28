import sys
import json
import os
import requests
from docx import Document
from docx.shared import Inches, Pt
import subprocess
from io import BytesIO
import re
import html

def strip_html(text):
    if not text:
        return ""
    text = text.replace('</p>', '\n').replace('<br>', '\n').replace('<br/>', '\n').replace('<br />', '\n')
    clean = re.compile('<.*?>')
    cleantext = re.sub(clean, '', text)
    return html.unescape(cleantext).strip()

def generate_surat_judul_pdf(json_path, template_path, output_pdf_path, logo_path):
    try:
        with open(json_path, 'r') as f:
            data = json.load(f)
        doc = Document(template_path)

        # 1. Fill basic placeholders in paragraphs
        for p in doc.paragraphs:
            if "{{NAMA}}" in p.text: p.text = p.text.replace("{{NAMA}}", str(data.get('studentName', '')))
            if "{{NIM}}" in p.text: p.text = p.text.replace("{{NIM}}", str(data.get('studentNim', '')))
            if "{{SEMESTER}}" in p.text: p.text = p.text.replace("{{SEMESTER}}", str(data.get('studentSemester', '-')))
            if "{{PRODI}}" in p.text: p.text = p.text.replace("{{PRODI}}", str(data.get('studentProdi', 'Sistem Informasi')))
            if "{{JUDUL_1}}" in p.text: p.text = p.text.replace("{{JUDUL_1}}", strip_html(data.get('judulPenelitian', '')))
            if "{{PEMBIMBING_1}}" in p.text:
                nama = str(data.get('dosenPaNama', '-'))
                nip = data.get('dosenPaNip')
                display = f"{nama} (NIP. {nip})" if nip and "..." not in nip else nama
                p.text = p.text.replace("{{PEMBIMBING_1}}", display)
            
            if "{{PEMBIMBING_2}}" in p.text:
                nama = str(data.get('pembimbing2Nama', '-'))
                nip = data.get('pembimbing2Nip')
                display = f"{nama} (NIP. {nip})" if nip and "..." not in nip else nama
                p.text = p.text.replace("{{PEMBIMBING_2}}", display)
            if "Palembang, ________________" in p.text:
                p.text = f"Palembang, {data.get('tanggal', '')}"

        # 2. Fill placeholders in tables (Signature Section and Header)
        for table in doc.tables:
            # Check if this is the header table
            is_header = any("UIN RADEN FATAH" in cell.text for row in table.rows for cell in row.cells)
            
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    # Header Logo insertion (Column 0, spans 2 rows)
                    if is_header and r_idx == 0 and c_idx == 0 and os.path.exists(logo_path):
                        # Clear cell and add picture
                        p = cell.paragraphs[0]
                        p.text = ""
                        p.alignment = 1 # Center
                        run = p.add_run()
                        run.add_picture(logo_path, width=Inches(0.8))
                    
                    for p in cell.paragraphs:
                        text = p.text
                        if "{{NAMA}}" in text: p.text = p.text.replace("{{NAMA}}", str(data.get('studentName', '')))
                        if "{{NIM}}" in text: p.text = p.text.replace("{{NIM}}", str(data.get('studentNim', '')))
                        if "{{KAPRODI_NAMA}}" in text: p.text = p.text.replace("{{KAPRODI_NAMA}}", str(data.get('kaprodiNama', '')))
                        if "{{KAPRODI_NIP}}" in text: p.text = p.text.replace("{{KAPRODI_NIP}}", str(data.get('kaprodiNip', '')))

        # Save temporary docx
        temp_docx = output_pdf_path.replace('.pdf', '.docx')
        doc.save(temp_docx)

        # Convert to PDF using LibreOffice
        subprocess.run([
            'libreoffice', '--headless', '--convert-to', 'pdf',
            '--outdir', os.path.dirname(output_pdf_path),
            temp_docx
        ], check=True)

        if os.path.exists(temp_docx):
            os.remove(temp_docx)

        print(f"PDF generated successfully at {output_pdf_path}")

    except Exception as e:
        print(f"Error: {str(e)}")
        sys.exit(1)

if __name__ == "__main__":
    if len(sys.argv) < 5:
        print("Usage: python3 generate_surat_judul_pdf.py <json_path> <template_path> <output_pdf_path> <logo_path>")
        sys.exit(1)
    
    generate_surat_judul_pdf(sys.argv[1], sys.argv[2], sys.argv[3], sys.argv[4])
