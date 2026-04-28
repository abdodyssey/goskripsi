import sys
import json
import docx
import os
import subprocess

def generate_undangan(config):
    exam_type = config.get('type')
    data = config.get('data', {})
    output_pdf = config.get('outputPath')
    
    template_map = {
        'sempro': 'Undangan Seminar Proposal.docx',
        'hasil': 'Undangan Ujian Hasil.docx',
        'skripsi': 'Undangan Ujian Skripsi.docx'
    }
    
    template_name = template_map.get(exam_type)
    if not template_name:
        raise ValueError(f"Unknown exam type: {exam_type}")
        
    template_path = os.path.join(os.path.dirname(__file__), '../templates/docs', template_name)
    doc = docx.Document(template_path)
    
    # Replacement logic
    def replace_text(content):
        for key, value in data.items():
            placeholder = f"{{{{{key}}}}}"
            if placeholder in content:
                content = content.replace(placeholder, str(value))
        return content

    # Replace in paragraphs
    for p in doc.paragraphs:
        p.text = replace_text(p.text)
        
    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.text = replace_text(p.text)
                    
    # Save temp docx
    temp_docx = output_pdf.replace('.pdf', '.docx')
    doc.save(temp_docx)
    
    # Convert to PDF using LibreOffice
    try:
        subprocess.run([
            'libreoffice', '--headless', '--convert-to', 'pdf',
            temp_docx, '--outdir', os.path.dirname(output_pdf)
        ], check=True)
        
        # Cleanup temp docx
        if os.path.exists(temp_docx):
            os.remove(temp_docx)
            
        print(f"Successfully generated {output_pdf}")
    except Exception as e:
        print(f"Error converting to PDF: {e}")
        sys.exit(1)

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python generate_undangan_ujian_pdf.py <json_config_path>")
        sys.exit(1)
        
    with open(sys.argv[1], 'r') as f:
        config = json.load(f)
        
    generate_undangan(config)
